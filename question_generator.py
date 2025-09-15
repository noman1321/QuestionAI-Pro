"""
Question Generator Module - Fixed Version
Handles text extraction, question generation, and formatting
"""

import streamlit as st
from openai import OpenAI
import PyPDF2
import docx
import requests
from bs4 import BeautifulSoup
import os
from dotenv import load_dotenv
import time
import re
from image_processor import ImageProcessor
from openai import OpenAI

# Load local .env if available
load_dotenv()

def get_openai_api_key():
    # First check Streamlit secrets, then .env
    return st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")

def get_client():
    api_key = get_openai_api_key()
    if not api_key:
        raise ValueError("OpenAI API key not found. Please set it in Streamlit Secrets or .env")
    return OpenAI(api_key=api_key)

# Load environment variables
# load_dotenv()

# def get_openai_api_key():
#     """Get OpenAI API key from multiple sources"""
#     try:
#         return st.secrets["OPENAI_API_KEY"]
#     except:
#         pass
#     return os.getenv('OPENAI_API_KEY')

class QuestionGenerator:
    def __init__(self):
        api_key = get_openai_api_key()
        if not api_key:
            st.error("❌ OpenAI API key not found!")
            st.stop()
        self.client = OpenAI(api_key=api_key)
        self.image_processor = ImageProcessor()
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF file with better error handling"""
        try:
            # Reset file pointer to beginning
            pdf_file.seek(0)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            st.info(f"📖 PDF has {len(reader.pages)} pages")
            
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text += page_text + "\n"
                        st.info(f"Page {i+1}: {len(page_text)} characters extracted")
                    else:
                        st.warning(f"Page {i+1}: No text found (might be image-based)")
                except Exception as page_error:
                    st.warning(f"Could not extract text from page {i+1}: {str(page_error)}")
                    continue
            
            if not text.strip():
                st.warning("⚠️ PDF appears to be image-based or encrypted. No text extracted.")
                return ""
            
            st.success(f"✅ PDF: Extracted {len(text)} characters total")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file with better error handling"""
        try:
            # Reset file pointer
            docx_file.seek(0)
            doc = docx.Document(docx_file)
            text = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Also extract text from tables
            table_count = 0
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text += cell.text + "\t"
                    table_text += "\n"
                
                if table_text.strip():
                    text += f"\n[TABLE {table_count + 1}]\n{table_text}\n"
                    table_count += 1
            
            if not text.strip():
                st.warning("⚠️ DOCX file appears to be empty or contains no readable text.")
                return ""
            
            st.success(f"✅ DOCX: Extracted {len(text)} characters from {paragraph_count} paragraphs and {table_count} tables")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading DOCX: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, txt_file):
        """Extract text from TXT file with multiple encoding support"""
        try:
            # Reset file pointer
            txt_file.seek(0)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252', 'ascii']
            text = ""
            successful_encoding = None
            
            for encoding in encodings:
                try:
                    txt_file.seek(0)
                    text = txt_file.read().decode(encoding)
                    successful_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    st.warning(f"Error with encoding {encoding}: {str(e)}")
                    continue
            
            if not text.strip():
                st.warning("⚠️ TXT file appears to be empty or could not be decoded.")
                return ""
            
            st.success(f"✅ TXT: Extracted {len(text)} characters using {successful_encoding} encoding")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading TXT: {str(e)}")
            return ""
    
    def extract_text_from_url(self, url):
        """Extract text from URL with improved content extraction"""
        try:
            st.info(f"🌐 Fetching content from: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
            }
            
            response = requests.get(url, timeout=15, headers=headers)
            response.raise_for_status()  # Raise an error for bad status codes
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(["script", "style", "nav", "header", "footer", "aside", "iframe", "noscript"]):
                element.extract()
            
            # Try to find main content areas in order of preference
            content_selectors = [
                'main',
                'article',
                '[role="main"]',
                '.content',
                '.main-content',
                '.article-content',
                '.post-content',
                '#content',
                '#main'
            ]
            
            main_content = None
            for selector in content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            if main_content:
                text = main_content.get_text()
            else:
                # Fallback to body content
                body = soup.find('body')
                text = body.get_text() if body else soup.get_text()
            
            # Clean up the text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            if not text.strip():
                st.warning(f"⚠️ No readable content found at {url}")
                return ""
            
            st.success(f"✅ URL: Extracted {len(text)} characters")
            return text
            
        except requests.exceptions.Timeout:
            st.error(f"❌ Timeout while fetching {url}")
            return ""
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Error fetching URL {url}: {str(e)}")
            return ""
        except Exception as e:
            st.error(f"❌ Error extracting text from URL {url}: {str(e)}")
            return ""
    
    def clean_question_format(self, questions_text):
        """Clean and format the generated questions with better processing"""
        lines = questions_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive formatting
            line = re.sub(r'\*+', '', line)  # Remove asterisks
            line = re.sub(r'#+', '', line)   # Remove hash marks
            
            # Skip metadata lines
            skip_patterns = [
                r'Question number:',
                r'Question type:',
                r'Difficulty:',
                r'Points:',
                r'Time:',
                r'Marks:'
            ]
            
            if any(re.search(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                continue
            
            # Clean up the line and add if not empty
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        # Join lines and add proper spacing
        result = '\n'.join(cleaned_lines)
        
        # Add proper spacing between questions
        result = re.sub(r'\n([A-Z]?\d+\.)', r'\n\n\1', result)
        result = re.sub(r'\n(Section [A-Z]:)', r'\n\n\1', result)
        
        # Remove excessive line breaks
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result.strip()

    def format_questions_with_images(self, questions_text, image_analyses):
        """Parse questions and embed image references with improved logic"""
        if not image_analyses:
            return questions_text, []
        
        formatted_questions = []
        question_images = []  # Store which images are used in questions
        
        # Split by double newlines to get question blocks
        blocks = re.split(r'\n\n+', questions_text)
        
        for block in blocks:
            if not block.strip():
                continue
                
            # Check if this block contains a question
            if re.search(r'^[A-Z]?\d+\.', block.strip(), re.MULTILINE):
                formatted_q, used_images = self._process_question_with_images(block, image_analyses)
                formatted_questions.append(formatted_q)
                question_images.extend(used_images)
            else:
                # This is likely a section header or other content
                formatted_questions.append(block)
        
        return '\n\n'.join(formatted_questions), question_images
    
    def _process_question_with_images(self, question_text, image_analyses):
        """Process individual question to embed images with better pattern matching"""
        used_images = []
        
        # Find image references with multiple possible patterns
        image_patterns = [
            r'\[IMAGE:img_(\d+)\]',
            r'\[img_(\d+)\]',
            r'IMAGE\s*(\d+)',
            r'\(Image\s*(\d+)\)',
            r'<img_(\d+)>'
        ]
        
        all_matches = []
        for pattern in image_patterns:
            matches = re.findall(pattern, question_text, re.IGNORECASE)
            all_matches.extend(matches)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_matches = []
        for match in all_matches:
            if match not in seen:
                seen.add(match)
                unique_matches.append(match)
        
        for match in unique_matches:
            img_num = int(match) - 1  # Convert to 0-based index
            if 0 <= img_num < len(image_analyses):
                used_images.append({
                    'index': img_num,
                    'image': image_analyses[img_num]['image'],
                    'source': image_analyses[img_num]['source'],
                    'analysis': image_analyses[img_num]['analysis']
                })
        
        return question_text, used_images

    def generate_single_question_set(self, content, difficulty, num_questions, question_types, reference_style, set_number, image_analyses, temperature=0.7):
        """Generate a single question set with improved prompting"""
        difficulty_prompts = {
            "Easy": "Create basic, straightforward questions that test fundamental understanding and recall of key concepts. Focus on definitions, simple facts, and basic comprehension.",
            "Medium": "Create moderate difficulty questions that require analysis, application of concepts, and some critical thinking. Include problem-solving and interpretation tasks.",
            "Hard": "Create challenging questions that require synthesis, evaluation, critical analysis, and deep understanding of complex concepts. Focus on advanced application and creative thinking."
        }

        type_list = ", ".join(question_types)

        style_instruction = f"""
        Please use the reference question paper style shown below for formatting and tone, 
        but do NOT copy questions directly. Create new questions with similar structure and difficulty level.
        
        Reference Question Paper Style:
        {reference_style[:3000]}
        """ if reference_style else ""

        # Enhanced variation prompts for different sets
        variation_prompts = [
            "Focus on theoretical concepts, definitions, and foundational knowledge. Emphasize understanding of core principles.",
            "Emphasize practical applications, real-world scenarios, and case studies. Focus on how concepts are applied in practice.",
            "Include analytical and problem-solving questions. Focus on breaking down complex problems and logical reasoning.",
            "Concentrate on comparative analysis, critical evaluation, and synthesis of different viewpoints or approaches.",
            "Focus on higher-order thinking skills, creative problem-solving, and innovative applications of concepts."
        ]

        variation_prompt = variation_prompts[(set_number - 1) % len(variation_prompts)]

        # Enhanced image content processing
        image_content = ""
        if image_analyses:
            image_content = "\n\nAVAILABLE IMAGES FOR QUESTION CREATION:\n"
            for i, img_analysis in enumerate(image_analyses):
                image_content += f"\nImage {i+1} (Reference: [IMAGE:img_{i+1}]) from {img_analysis['source']}:\n"
                image_content += f"Analysis: {img_analysis['analysis'][:800]}{'...' if len(img_analysis['analysis']) > 800 else ''}\n"
            
            image_content += """\n
IMPORTANT INSTRUCTIONS FOR IMAGE-BASED QUESTIONS:
- Use the format [IMAGE:img_X] where X is the image number to reference images
- Create 2-4 questions that specifically relate to the images provided
- Distribute image questions across both sections appropriately
- Ensure image-based questions test understanding of the visual content
- Make questions specific to what's shown in the images
"""

        # Enhanced prompt with better structure
        prompt = f"""
You are creating a comprehensive examination question paper with the following structure:

**QUESTION PAPER STRUCTURE:**

**Section A: Compulsory Questions (10 Questions)**
- Mix of question types: {', '.join(t for t in question_types if t != 'Long Answer') if any(t != 'Long Answer' for t in question_types) else 'Multiple Choice Questions (MCQ), Short Answer, Fill in the Blanks, True/False'}
- Each question worth 2-3 marks
- Must include EXACTLY 10 questions

**Section B: Long Answer Questions (Choose any 7 out of 15)**
- Generate EXACTLY 15 long answer questions
- Each question worth 10-15 marks
- Students must attempt any 7 questions
- Include instruction: "Attempt any SEVEN questions from this section"

**QUESTION SET SPECIFICATIONS:**
- Difficulty Level: {difficulty}
- Difficulty Guidelines: {difficulty_prompts[difficulty]}
- Set {set_number} Focus: {variation_prompt}
- Total Questions Required: 25 (10 compulsory + 15 optional)

{style_instruction}

**SOURCE CONTENT:**
{content[:7000]}

{image_content}

**FORMATTING REQUIREMENTS:**
1. Use clear section headers: "**Section A: Compulsory Questions**" and "**Section B: Long Answer Questions**"
2. Number questions as A1, A2, ..., A10 and B1, B2, ..., B15
3. For MCQs: Use (A), (B), (C), (D) format with correct answer marked
4. For image questions: Place [IMAGE:img_X] reference at the start of the question
5. Use proper spacing between questions
6. No asterisks or excessive formatting
7. Include question type indicators where helpful

**EXAMPLE FORMAT:**
```
**Section A: Compulsory Questions** (Attempt ALL questions)

A1. [MCQ] What is the primary function of...?
(A) Option 1
(B) Option 2
(C) Option 3
(D) Option 4
Correct Answer: (B)

A2. [IMAGE:img_1] [Short Answer]
Based on the diagram shown above, explain the relationship between...

**Section B: Long Answer Questions** (Attempt any SEVEN questions)

B1. Discuss in detail the implications of... (15 marks)
```

CRITICAL REQUIREMENTS:
- Create completely UNIQUE questions different from other question sets
- Ensure exactly 10 questions in Section A and 15 questions in Section B
- {f"Include {min(4, len(image_analyses))} image-based questions using [IMAGE:img_X] format" if image_analyses else "Focus on text-based content only"}
- Vary question complexity and focus areas within the specified difficulty level
- Ensure questions test different aspects of the content provided

Generate the complete question paper now:
"""

        try:
            # Use the best available model
            models_to_try = ["gpt-4o", "gpt-4-turbo-preview", "gpt-4-turbo", "gpt-4"]
            
            for model in models_to_try:
                try:
                    response = self.client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are an expert educator and examination specialist who creates high-quality, unique examination questions. You can analyze images and create comprehensive question papers with proper structure and formatting."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=4000,
                        temperature=temperature
                    )
                    
                    result = response.choices[0].message.content
                    st.success(f"✅ Question set generated using {model}")
                    break
                    
                except Exception as model_error:
                    st.warning(f"⚠️ Model {model} failed: {str(model_error)}")
                    continue
            else:
                st.error("❌ All models failed to generate questions")
                return ""
            
            # Clean and format the questions
            cleaned_result = self.clean_question_format(result)
            
            return cleaned_result
            
        except Exception as e:
            st.error(f"❌ Error generating Set {set_number}: {str(e)}")
            return ""

    def generate_multiple_question_sets(self, content, difficulty, num_questions, question_types, reference_style, image_analyses, num_sets=5):
        """Generate multiple unique question sets with embedded images"""
        question_sets = []
        temperatures = [0.7, 0.8, 0.9, 0.6, 0.75]  # Different temperatures for variety
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i in range(num_sets):
            status_text.info(f"🤖 Generating Question Set {i+1} of {num_sets}...")
            
            questions = self.generate_single_question_set(
                content, difficulty, num_questions, question_types, 
                reference_style, i+1, image_analyses, temperatures[i % len(temperatures)]
            )
            
            if questions and questions.strip():
                # Format questions with embedded images
                formatted_questions, used_images = self.format_questions_with_images(questions, image_analyses)
                
                question_sets.append({
                    'set_number': i+1,
                    'questions': formatted_questions,
                    'used_images': used_images,  # Store which images are used
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
                    'word_count': len(formatted_questions.split()),
                    'character_count': len(formatted_questions)
                })
                st.success(f"✅ Question Set {i+1} generated successfully! ({len(formatted_questions.split())} words)")
            else:
                st.error(f"❌ Failed to generate Set {i+1}")
            
            progress_bar.progress((i+1) / num_sets)
            
            # Rate limiting with exponential backoff
            if i < num_sets - 1:
                delay = min(2 + (i * 0.5), 5)  # Increase delay slightly for each set
                time.sleep(delay)
        
        progress_bar.empty()
        status_text.empty()
        
        # Display generation summary
        if question_sets:
            st.success(f"🎉 Successfully generated {len(question_sets)} question sets!")
            total_words = sum(set_data['word_count'] for set_data in question_sets)
            st.info(f"📊 Total content generated: {total_words} words across all sets")
        

        return question_sets
